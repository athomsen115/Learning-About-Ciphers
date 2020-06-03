import sys, docx
from docx.shared import RGBColor, Pt

#This program uses the formatting in docx and steganography to hide a secret message in the blank lines of a fake message

fakeLetter = docx.Document('fakeMessage.docx')
fakeList = []
for paragraph in fakeLetter.paragraphs:
    fakeList.append(paragraph.text)

#realLetter = docx.Document('realMessage.docx')
realLetter = docx.Document('realMessage_vig.docx') #KEY is BOMBINGS
realList = []
for paragraph in realLetter.paragraphs:
    if len(paragraph.text) != 0:
        realList.append(paragraph.text)
        
def lineLimit(fake, real):
    numBlanks = 0
    numReal = 0
    for line in fake:
        if line == '':
            numBlanks += 1
    numReal = len(real)
    diff = numReal - numBlanks
    print("Number of blank lines in fake message: {}".format(numBlanks))
    print("Number of lines in real message: {}".format(numReal))
    if numReal > numBlanks:
        print("Fake message needs {} more blank lines".format(diff))
        sys.exit()
        
lineLimit(fakeList, realList)
        
template = docx.Document('template.docx')
template.add_heading('Winston Churchill', 0)
subtitle = template.add_heading('Prime Minister of Britain', 1)
subtitle.alignment = 1
template.add_heading('', 1)
template.add_paragraph('July 18, 1940')
template.add_paragraph('')

def setSpacing(paragraph):
    format = paragraph.paragraph_format
    format.space_before = Pt(0)
    format.space_after = Pt(0)

length = len(realList)
count = 0

for line in fakeList:
    if count < length and line == "":
        paragraph = template.add_paragraph(realList[count])
        paragraphIndex = len(template.paragraphs) - 1
        run = template.paragraphs[paragraphIndex].runs[0]
        font = run.font
        font.color.rgb = RGBColor(255,0,0)
        count +=1
    else: 
        paragraph = template.add_paragraph(line)
        
    setSpacing(paragraph)

template.save('ciphertext.docx') 
print("Done")
